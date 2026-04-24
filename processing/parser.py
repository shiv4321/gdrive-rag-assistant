"""
processing/parser.py
──────────────────────────────────────────────────────────────────
Stateless text extraction from raw bytes.

  extract_text(data: bytes, mime_type: str) → str

Supported formats:
  • application/pdf                                     → PyMuPDF
  • application/vnd.openxmlformats-officedocument…docx → python-docx
  • text/plain                                          → UTF-8 decode
  • application/vnd.google-apps.document               → exported as docx
"""

from __future__ import annotations

import io
import logging
import re

import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)


def extract_text(data: bytes, mime_type: str) -> str:
    """
    Extract plain text from raw file bytes.

    Parameters
    ----------
    data      : raw bytes of the file
    mime_type : original MIME type (used to pick the parser)

    Returns
    -------
    Cleaned, normalised plain text string.
    """
    if mime_type == "application/pdf":
        return _parse_pdf(data)

    if mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/vnd.google-apps.document",   # exported to docx upstream
    ):
        return _parse_docx(data)

    if mime_type == "text/plain":
        return _parse_txt(data)

    # Fallback: try UTF-8 then latin-1
    logger.warning("Unknown mime_type=%s – attempting plain-text fallback", mime_type)
    return _parse_txt(data)


def _parse_pdf(data: bytes) -> str:
    text_parts: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text("text"))
    return _clean("\n".join(text_parts))


def _parse_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return _clean("\n".join(paragraphs))


def _parse_txt(data: bytes) -> str:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1", errors="replace")
    return _clean(text)


def _clean(text: str) -> str:
    """Normalise whitespace, remove null bytes and control characters."""
    text = text.replace("\x00", "")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(line.rstrip() for line in text.splitlines())
    return text.strip()
